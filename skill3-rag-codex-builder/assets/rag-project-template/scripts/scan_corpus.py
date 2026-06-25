#!/usr/bin/env python3
"""Extract text from local documents into JSONL for a RAG pipeline."""

from __future__ import annotations

import argparse
import hashlib
import json
import zipfile
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree


SUPPORTED = {".txt", ".md", ".markdown", ".docx", ".pdf"}


def read_docx_with_stdlib(path: Path) -> str:
    """Fallback DOCX reader for environments without python-docx."""
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml")
    root = ElementTree.fromstring(xml)
    parts = []
    for paragraph in root.findall(".//w:p", namespace):
        texts = [node.text or "" for node in paragraph.findall(".//w:t", namespace)]
        text = "".join(texts).strip()
        if text:
            parts.append(text)
    return "\n".join(parts)


def read_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        try:
            from docx import Document
        except Exception:  # pragma: no cover
            return read_docx_with_stdlib(path)
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except Exception as exc:  # pragma: no cover
            raise RuntimeError("Install pypdf to read PDF files") from exc
        reader = PdfReader(str(path))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    return ""


def iter_files(root: Path) -> Iterable[Path]:
    for path in sorted(root.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED:
            yield path


def doc_id(path: Path, root: Path) -> str:
    rel = path.relative_to(root).as_posix()
    digest = hashlib.sha1(rel.encode("utf-8")).hexdigest()[:10]
    return f"{path.stem}-{digest}"


def main() -> None:
    parser = argparse.ArgumentParser(description="Scan local documents into corpus.jsonl")
    parser.add_argument("source_dir", help="Directory containing source documents")
    parser.add_argument("--out", default="build/corpus.jsonl", help="Output JSONL path")
    args = parser.parse_args()

    source = Path(args.source_dir).resolve()
    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)

    count = 0
    with out.open("w", encoding="utf-8") as handle:
        for path in iter_files(source):
            text = read_text(path).strip()
            if not text:
                continue
            record = {
                "doc_id": doc_id(path, source),
                "title": path.stem,
                "source_path": str(path),
                "modified_time": path.stat().st_mtime,
                "text": text,
            }
            handle.write(json.dumps(record, ensure_ascii=False) + "\n")
            count += 1

    print(json.dumps({"output": str(out), "documents": count}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
